"""法律文书生成工具"""
from typing import Dict, Any, Optional
import os
from pathlib import Path
from datetime import datetime
from .base import BaseTool

# 处理相对导入问题
try:
    from ..config.config import Config
except (ImportError, ValueError):
    import sys
    from pathlib import Path as PathLib
    current_file = PathLib(__file__).resolve()
    project_root = current_file.parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))
    from config.config import Config


class DocumentGeneratorTool(BaseTool):
    """法律文书生成工具，用于生成docx或markdown格式的法律文书"""
    
    def __init__(self, config: Optional[Config] = None):
        """
        初始化文书生成工具
        
        Args:
            config: 系统配置
        """
        super().__init__(
            name="generate_legal_document",
            description="用于起草、生成法律文书或合同。当用户请求写一份文件（如起诉状、离婚协议、合同等）时调用此工具。支持生成docx和markdown格式的文件。"
        )
        self.config = config or Config()
        
        # 创建输出目录
        self.output_dir = Path("./output")
        self.output_dir.mkdir(exist_ok=True)
    
    def execute(self, user_input: str, context: Dict[str, Any] = None) -> str:
        """
        执行文书生成
        
        Args:
            user_input: 用户输入（JSON格式字符串，包含title, content, file_format）
            context: 上下文信息（可选，可能包含title, content, file_format等）
            
        Returns:
            生成文件的绝对路径字符串，格式："文件已生成: /abs/path/to/file.docx"
        """
        try:
            # 解析输入参数
            # 优先从context获取，然后尝试从user_input解析JSON
            if context:
                title = context.get("title") or self._extract_from_json(user_input, "title")
                content = context.get("content") or self._extract_from_json(user_input, "content")
                file_format = context.get("file_format") or self._extract_from_json(user_input, "file_format", "docx")
            else:
                # 尝试从user_input解析JSON
                title = self._extract_from_json(user_input, "title")
                content = self._extract_from_json(user_input, "content")
                file_format = self._extract_from_json(user_input, "file_format", "docx")
            
            # 验证必需参数
            if not title:
                return "Error: 缺少必需参数 'title'（文书标题）"
            if not content:
                return "Error: 缺少必需参数 'content'（文书内容）"
            
            # 验证文件格式
            if file_format not in ["docx", "markdown", "md"]:
                file_format = "docx"  # 默认使用docx
            
            # 生成文件名（使用时间戳避免冲突）
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
            safe_title = safe_title.replace(' ', '_')
            
            if file_format == "markdown" or file_format == "md":
                filename = f"{safe_title}_{timestamp}.md"
                file_path = self.output_dir / filename
                self._generate_markdown(file_path, title, content)
            else:
                filename = f"{safe_title}_{timestamp}.docx"
                file_path = self.output_dir / filename
                self._generate_docx(file_path, title, content)
            
            # 返回绝对路径
            abs_path = file_path.resolve()
            return f"文件已生成: {abs_path}"
            
        except Exception as e:
            return f"Error: 生成文书时出错 - {str(e)}"
    
    def _extract_from_json(self, text: str, key: str, default: Any = None) -> Any:
        """从JSON字符串中提取值"""
        try:
            import json
            # 尝试解析JSON
            if text.strip().startswith('{'):
                data = json.loads(text)
                return data.get(key, default)
        except:
            pass
        return default
    
    def _generate_markdown(self, file_path: Path, title: str, content: str):
        """生成Markdown格式文件"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")
            f.write(content)
    
    def _generate_docx(self, file_path: Path, title: str, content: str):
        """生成DOCX格式文件"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")
        
        # 创建文档
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        time_para = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para_format = time_para.runs[0].font
        time_para_format.size = Pt(10)
        time_para_format.color = RGBColor(128, 128, 128)
        
        # 添加分隔线
        doc.add_paragraph("─" * 50)
        
        # 添加内容（按段落分割）
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                # 检查是否是标题（以#开头）
                if para_text.strip().startswith('#'):
                    level = len(para_text) - len(para_text.lstrip('#'))
                    heading_text = para_text.lstrip('#').strip()
                    doc.add_heading(heading_text, level=min(level, 3))
                else:
                    doc.add_paragraph(para_text)
            else:
                doc.add_paragraph()  # 空行
        
        # 保存文档
        doc.save(str(file_path))
    
    def to_schema(self) -> Dict[str, Any]:
        """转换为OpenAI格式的JSON Schema"""
        return {
            "type": "function",
            "function": {
                "name": self.name,
                "description": self.description,
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "文书的标题，例如：'离婚协议书'、'民事起诉状'、'劳动合同'等"
                        },
                        "content": {
                            "type": "string",
                            "description": "文书的正文内容，应包含完整的法律文书内容，可以使用Markdown格式（标题、段落、列表等）"
                        },
                        "file_format": {
                            "type": "string",
                            "description": "文件格式，可选值：'docx'（Word文档）或'markdown'/'md'（Markdown文档）",
                            "enum": ["docx", "markdown", "md"],
                            "default": "docx"
                        }
                    },
                    "required": ["title", "content"]
                }
            }
        }

